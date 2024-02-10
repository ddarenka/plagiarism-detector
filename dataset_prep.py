import docx
import pandas as pd

def extract_text_from_docx(docx_path): # extract text
    text = ""
    doc = docx.Document(docx_path)
    for paragraph in doc.paragraphs:
        if paragraph.text: 
            text += paragraph.text + " "  
    return text.strip() 


def extract_data_from_table(docx_path, marker): # create contents
    doc = docx.Document(docx_path)
    table_data = []
    
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)

    # remove elements after marker
    try:
        index_to_remove = table_data.index(marker)
        table_data[index_to_remove:] = []
    except ValueError:
        pass

    return table_data

# for satellite 2020:
def extract_table_of_contents(docx_path, markers):
    doc = docx.Document(docx_path)
    toc_data = []

    found_markers = set(markers)
    found_marker = False

    for paragraph in doc.paragraphs:
        if found_marker:
            toc_data.append(paragraph.text)
        else:
            for marker in found_markers:
                if marker in paragraph.text:
                    found_marker = True
                    break

    return toc_data

def process_contents_df(contents_data):
    df = pd.DataFrame(contents_data, columns=['Article_Name', 'Page'])

    # split 'Article_Name' into 'Authors' and 'Article_Name'
    df['Authors'], df['Article_Name'] = df['Article_Name'].str.rsplit('.', 1).str
    # clean spaces
    df['Article_Name'] = df['Article_Name'].str.strip()
    # add period
    df['Authors'] = df['Authors'] + '.'
    # remove rows where 'Article_Name' is NaN
    df = df.dropna(subset=['Article_Name'])
    # replace '\n' with spaces
    df['Article_Name'] = df['Article_Name'].replace('\n', ' ', regex=True)
    # reset index
    df.reset_index(drop=True, inplace=True)

    return df


def find_article_text(text, start, end=None):
    start_index = text.find(start)
    if start_index == -1:
        return None # no article found

    start_index += len(start)
    
    if end:
        end_index = text.find(end, start_index)
        if end_index == -1:
            return start_index + 'ERROR' # no end of the article found
        return text[start_index:end_index]
    else:
        return text[start_index:]
    

def articles_text_to_df(contents_df, text_from_docx, end_marker='References:'):
    df = contents_df.copy()
    df['Text'] = None

    for i in range(len(df)):
        start = df['Article_Name'][i].upper()
        df.at[i, 'Text'] = find_article_text(text_from_docx, start, end_marker)

    return df


def create_total_dataframe(all_docx):
    total_df = pd.DataFrame()

    for docx_path in all_docx:
        text_from_docx = extract_text_from_docx(docx_path)

        if '2022' in docx_path:
            marker = ['Authors', '180']
        elif '2021' in docx_path:
            marker = ['Figure 1: Centralized network', 'Figure 2: Decentralized network']
        else:
            marker = None

        contents_data = extract_data_from_table(docx_path, marker)
        contents_df = process_contents_df(contents_data)
        articles_df = articles_text_to_df(contents_df, text_from_docx)

        total_df = pd.concat([total_df, articles_df], ignore_index=True)

    return total_df